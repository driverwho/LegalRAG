"""DOCX loader with OCR support for image extraction."""

import logging
from typing import List

import numpy as np
import tqdm
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader
from PIL import Image

from backend.app.core.document.ocr import get_ocr

logger = logging.getLogger(__name__)


class RapidOCRDocLoader(UnstructuredFileLoader):
    """Word document loader that extracts text from both text content and images using OCR.

    Extends UnstructuredFileLoader to provide enhanced DOCX loading with
    image text extraction via RapidOCR.
    """

    def _get_elements(self) -> List:
        """Extract text elements from DOCX including OCR from embedded images."""

        def doc2text(filepath: str) -> str:
            """Extract text from DOCX including OCR from embedded images.

            Args:
                filepath: Path to DOCX file

            Returns:
                Combined text content from document text and image OCR
            """
            from io import BytesIO

            from docx import Document, ImagePart
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table, _Cell
            from docx.text.paragraph import Paragraph

            ocr = RapidOCR()
            doc = Document(filepath)
            resp = ""

            def iter_block_items(parent):
                """Iterate over document blocks (paragraphs and tables)."""
                from docx.document import Document as DocType

                if isinstance(parent, DocType):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(
                total=len(doc.paragraphs) + len(doc.tables),
                desc="RapidOCRDocLoader block index: 0",
            )
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description("RapidOCRDocLoader block index: {}".format(i))
                b_unit.refresh()

                if isinstance(block, Paragraph):
                    # Extract paragraph text
                    resp += block.text.strip() + "\n"

                    # Extract and OCR images from paragraph
                    images = block._element.xpath(".//pic:pic")  # Get all images
                    for image in images:
                        for img_id in image.xpath(".//a:blip/@r:embed"):  # Get image id
                            try:
                                part = doc.part.related_parts[img_id]  # Get image by id
                                if isinstance(part, ImagePart):
                                    img = Image.open(BytesIO(part._blob))
                                    # Convert to RGB if necessary
                                    if img.mode != "RGB":
                                        img = img.convert("RGB")
                                    result, _ = ocr(np.array(img))
                                    if result:
                                        ocr_result = [line[1] for line in result]
                                        resp += "\n".join(ocr_result) + "\n"
                            except Exception as img_exc:
                                logger.warning(
                                    f"Failed to process image in block {i}: {img_exc}"
                                )
                                continue

                elif isinstance(block, Table):
                    # Extract table cell text
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"

                b_unit.update(1)

            return resp

        text = doc2text(self.file_path)
        from unstructured.partition.text import partition_text

        return partition_text(text=text, **self.unstructured_kwargs)
